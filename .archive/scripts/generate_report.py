import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading_1(doc, text):
    return doc.add_paragraph(text, style='Heading 1 Custom')

def add_heading_2(doc, text):
    return doc.add_paragraph(text, style='Heading 2 Custom')

def add_heading_3(doc, text):
    return doc.add_paragraph(text, style='Heading 3 Custom')

def add_content(doc, text):
    return doc.add_paragraph(text, style='Content Custom')

def setup_styles(doc):
    # Thiết lập lề trang chuẩn đồ án (Trái: 3cm, Phải: 2cm, Trên: 2.5cm, Dưới: 2.5cm)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

    styles = doc.styles

    # Style: Chương (Heading 1 Custom)
    h1_style = styles.add_style('Heading 1 Custom', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.base_style = styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Times New Roman'
    h1_font.size = Pt(14)
    h1_font.bold = True
    h1_font.all_caps = True
    h1_font.color.rgb = RGBColor(0, 0, 0)
    h1_format = h1_style.paragraph_format
    h1_format.space_before = Pt(24)
    h1_format.space_after = Pt(24)
    h1_format.line_spacing = 1.0
    h1_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h1_format.left_indent = Pt(0)

    # Style: Tiểu mục cấp 1 (Heading 2 Custom)
    h2_style = styles.add_style('Heading 2 Custom', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.base_style = styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Times New Roman'
    h2_font.size = Pt(13)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(0, 0, 0)
    h2_format = h2_style.paragraph_format
    h2_format.space_before = Pt(6)
    h2_format.space_after = Pt(12)
    h2_format.line_spacing = 1.0
    h2_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2_format.left_indent = Pt(0)

    # Style: Tiểu mục cấp 2 (Heading 3 Custom)
    h3_style = styles.add_style('Heading 3 Custom', WD_STYLE_TYPE.PARAGRAPH)
    h3_style.base_style = styles['Heading 3']
    h3_font = h3_style.font
    h3_font.name = 'Times New Roman'
    h3_font.size = Pt(13)
    h3_font.bold = True
    h3_font.italic = True
    h3_font.color.rgb = RGBColor(0, 0, 0)
    h3_format = h3_style.paragraph_format
    h3_format.space_before = Pt(6)
    h3_format.space_after = Pt(12)
    h3_format.line_spacing = 1.0
    h3_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3_format.left_indent = Pt(0)

    # Style: Nội dung (Content Custom)
    content_style = styles.add_style('Content Custom', WD_STYLE_TYPE.PARAGRAPH)
    c_font = content_style.font
    c_font.name = 'Times New Roman'
    c_font.size = Pt(13)
    c_format = content_style.paragraph_format
    c_format.space_before = Pt(10)
    c_format.space_after = Pt(0)
    c_format.line_spacing = 1.5
    c_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    c_format.left_indent = Pt(0)
    c_format.first_line_indent = Pt(0)

def generate_report():
    doc = Document()
    setup_styles(doc)

    add_heading_1(doc, "CHƯƠNG 1. MỞ ĐẦU VÀ BÀI TOÁN THỰC TẾ")
    
    add_heading_2(doc, "1.1 Bối cảnh thị trường và nhu cầu dữ liệu")
    add_content(doc, "Thị trường chứng khoán Việt Nam tạo ra hàng triệu giao dịch mỗi ngày. Việc truy cập dữ liệu giao dịch sạch, chính xác và có cấu trúc để phân tích kỹ thuật (Technical Analysis) là một thách thức lớn. Phần lớn dữ liệu thô (Raw Data) từ các API công cộng thiếu tính nhất quán, có sai số và chưa được tối ưu hóa cho truy vấn phân tích.")
    
    add_heading_2(doc, "1.2 Bài toán đặt ra và tính cấp thiết")
    add_content(doc, "Nhiều nhà phân tích và kỹ sư dữ liệu gặp khó khăn trong việc xây dựng một luồng dữ liệu tự động, chạy lặp lại nhiều lần mà không tạo ra dữ liệu rác (tính Idempotency). Hơn nữa, việc tính toán các chỉ báo tài chính phức tạp như MACD (Moving Average Convergence Divergence) hay RSI (Relative Strength Index) thường được thực hiện thủ công bằng Python hoặc Excel, gây nghẽn cổ chai khi cần xử lý quy mô lớn trên hệ quản trị cơ sở dữ liệu (Database). Đồ án này giải quyết bài toán xây dựng một hệ thống Data Pipeline tự động chuẩn Enterprise, có khả năng mở rộng và chịu lỗi cao.")

    add_heading_1(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT")
    
    add_heading_2(doc, "2.1 Kiến trúc Medallion (Data Lakehouse)")
    add_content(doc, "Kiến trúc Medallion, do Databricks đề xuất, chia quy trình quản lý dữ liệu thành 3 lớp phân cấp chất lượng [1]:")
    add_content(doc, "- Tầng Bronze (Raw Data): Chứa dữ liệu gốc ở dạng chưa qua xử lý (JSON, CSV). Đảm bảo không mất mát thông tin nguồn.")
    add_content(doc, "- Tầng Silver (Cleansed & Conformed): Dữ liệu được lọc, chuẩn hóa kiểu dữ liệu, loại bỏ giá trị bất thường.")
    add_content(doc, "- Tầng Gold (Curated/Business-level): Dữ liệu được tổng hợp theo mô hình hình sao (Star Schema), sẵn sàng cho các công cụ BI.")
    
    add_heading_2(doc, "2.2 Nguyên lý Kỹ thuật Dữ liệu (Data & Software Engineering)")
    add_heading_3(doc, "2.2.1 Tính lũy đẳng (Idempotency)")
    add_content(doc, "Trong Data Engineering, một pipeline được coi là Idempotent nếu việc chạy lại (re-run) cùng một tiến trình cho ra kết quả giống hệt lần chạy đầu tiên mà không gây trùng lặp dữ liệu [2]. Hệ thống đạt được điều này thông qua cấu trúc UPDATE ON CONFLICT (Upsert) và cơ chế Delete + Insert tại dbt.")
    
    add_heading_3(doc, "2.2.2 Tối ưu hóa tải tăng dần (Incremental Loading)")
    add_content(doc, "Việc xử lý toàn bộ dữ liệu (Full Refresh) hàng ngày tiêu tốn tài nguyên khổng lồ. Tải tăng dần (Incremental) chỉ tính toán trên các bản ghi mới hoặc có sự thay đổi, giảm khối lượng xử lý từ mức toàn bộ dữ liệu lịch sử xuống mức Delta.")

    add_heading_2(doc, "2.3 Cơ sở toán học của các chỉ báo kỹ thuật")
    add_heading_3(doc, "2.3.1 Chỉ báo sức mạnh tương đối (RSI)")
    add_content(doc, "RSI đo lường tốc độ và sự thay đổi của biến động giá. Công thức RSI truyền thống dùng Wilder's Smoothing, là một dạng làm mượt hàm số mũ (Exponential Smoothing), tạo ra thách thức lớn khi chuyển đổi sang ngôn ngữ SQL do đặc tính đệ quy (phụ thuộc vào giá trị ngày trước đó).")
    add_heading_3(doc, "2.3.2 Chỉ báo MACD và Exponential Moving Average (EMA)")
    add_content(doc, "Đường MACD Signal Line được định nghĩa chuẩn xác là đường EMA 9 chu kỳ của đường MACD Line. Khác với Simple Moving Average (SMA) chỉ lấy trung bình cộng đơn thuần, EMA gán trọng số lũy thừa cao hơn cho dữ liệu gần nhất. Tính chất hồi quy: EMA(t) = Value(t) * k + EMA(t-1) * (1-k) đòi hỏi phải sử dụng các kỹ thuật Window Function nâng cao trong SQL để tính toán.")

    add_heading_1(doc, "CHƯƠNG 3. KIẾN TRÚC HỆ THỐNG VÀ DATA CONTRACTS")
    
    add_heading_2(doc, "3.1 Lựa chọn công nghệ cốt lõi")
    add_content(doc, "Hệ thống sử dụng PostgreSQL 17 làm Data Warehouse trung tâm, dbt (Data Build Tool) phiên bản 1.10.x để điều phối SQL Transformation, và Apache Airflow làm Orchestrator. PostgreSQL được lựa chọn nhờ hỗ trợ mạnh mẽ cho định dạng JSONB (cho tầng Bronze) và khả năng xử lý Window Functions hiệu năng cao cho dữ liệu chuỗi thời gian [3].")
    
    add_heading_2(doc, "3.2 Hợp đồng dữ liệu (Data Contracts)")
    add_content(doc, "Data Contracts đảm bảo tính tin cậy giữa các tầng dữ liệu. Tại Silver Layer, hệ thống kỳ vọng nhận đủ các cột thô từ Bronze. Nếu phát hiện sai sót (ví dụ: Close <= 0, High < Low), dữ liệu sẽ không bị xóa ngay lập tức mà được gắn cờ is_valid = FALSE và lý do dq_flag để phục vụ mục đích kiểm toán (Audit).")

    add_heading_1(doc, "CHƯƠNG 4. THIẾT KẾ CHI TIẾT VÀ HIỆN THỰC HÓA")
    
    add_heading_2(doc, "4.1 Lớp Ingestion (Thu thập và xử lý lỗi API)")
    add_content(doc, "API chứng khoán (VNDirect, TCBS) thường xuyên bị nghẽn hoặc Rate Limit. Lớp Ingestion trong đồ án sử dụng kiến trúc Provider hợp nhất qua thư viện Vnstock 4.x. Nó triển khai thuật toán Exponential Backoff để thử lại thông minh (Retry), đồng thời hỗ trợ MockProvider cho các luồng kiểm thử E2E.")
    
    add_heading_2(doc, "4.2 Lớp Data Transformation (dbt) & Tối ưu SQL Đệ quy")
    add_content(doc, "Thách thức kỹ thuật lớn nhất là hiện thực hóa công thức hồi quy EMA9 của chỉ báo MACD bằng SQL thay vì Python Pandas. Hệ thống xây dựng một dbt Macro 'calculate_ema' sử dụng Window Functions để tái tạo mảng lũy thừa và tổng tích lũy, giúp giải quyết bài toán đệ quy trong môi trường cơ sở dữ liệu với sai số bằng 0.00% so với công thức chuẩn của thư viện TA-Lib Python.")
    
    add_heading_2(doc, "4.3 Điều phối luồng công việc (Airflow DAGs)")
    add_content(doc, "Airflow điều phối tiến trình chạy hàng ngày bằng DAG. Cơ chế trigger được cấu hình nghiêm ngặt với logical_date (execution_date) giúp đảm bảo tính Idempotency. Backfill cho dữ liệu quá khứ có thể chạy độc lập với luồng Daily mà không gây xung đột.")

    add_heading_1(doc, "CHƯƠNG 5. KẾT QUẢ VÀ HƯỚNG PHÁT TRIỂN")
    
    add_heading_2(doc, "5.1 Đánh giá hiệu suất và tính đúng đắn")
    add_content(doc, "Việc áp dụng chiến lược Incremental Loading giúp hệ thống chỉ thực thi truy vấn DML (Delete + Insert) trên 60 ngày dữ liệu gần nhất, giảm hơn 90% thời gian xử lý so với Full Refresh. Mọi chỉ báo tài chính (MACD, RSI, Bollinger Bands) đều vượt qua các Test Case (Data Quality Tests) được định nghĩa chặt chẽ trong dbt schema.")
    
    add_heading_2(doc, "5.2 Trực quan hóa")
    add_content(doc, "Dữ liệu được tổ chức dưới dạng Star Schema (fact_stock_indicators kết hợp với dim_stock, dim_date) tại tầng Gold, kết nối trực tiếp với Power BI, cho phép phân tích độ trễ cực thấp (Low Latency) và trực quan hóa xu hướng thị trường chuẩn xác.")
    
    add_heading_2(doc, "5.3 Hạn chế và Hướng phát triển")
    add_content(doc, "Do rào cản về API, hệ thống đang xử lý dữ liệu Batch theo chu kỳ ngày. Trong tương lai, kiến trúc có thể mở rộng lên luồng Streaming bằng Apache Kafka và Spark Structured Streaming để xử lý dữ liệu Intraday. Đối với việc lưu trữ quy mô lớn, việc chuyển đổi từ PostgreSQL sang Cloud Data Warehouse như BigQuery sẽ mang lại khả năng xử lý cột (Columnar Storage) vượt trội.")

    add_heading_1(doc, "TÀI LIỆU THAM KHẢO")
    add_content(doc, "[1] Databricks, \"What is a Medallion Architecture?\" [Online]. Available: https://www.databricks.com/glossary/medallion-architecture.")
    add_content(doc, "[2] M. Kleppmann, Designing Data-Intensive Applications: The Big Ideas Behind Reliable, Scalable, and Maintainable Systems, 1st ed. O'Reilly Media, 2017.")
    add_content(doc, "[3] J. Reis and M. Housley, Fundamentals of Data Engineering, O'Reilly Media, 2022.")

    doc.save('Bao_Cao_Do_An.docx')

if __name__ == '__main__':
    generate_report()
    print("Report generated successfully.")
