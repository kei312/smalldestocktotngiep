import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading_1(doc, text):
    p = doc.add_paragraph(text, style='Heading 1 Custom')
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph(text, style='Heading 2 Custom')
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph(text, style='Heading 3 Custom')
    return p

def add_content(doc, text):
    p = doc.add_paragraph(text, style='Content Custom')
    return p

def setup_styles(doc):
    # Setup Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

    styles = doc.styles

    # Heading 1 Custom
    h1_style = styles.add_style('Heading 1 Custom', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.base_style = styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Times New Roman'
    h1_font.size = Pt(14)
    h1_font.bold = True
    h1_font.all_caps = True
    h1_font.color.rgb = RGBColor(0, 0, 0)
    h1_format = h1_style.paragraph_format
    h1_format.space_before = Pt(24)
    h1_format.space_after = Pt(24)
    h1_format.line_spacing = 1.0
    h1_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h1_format.left_indent = Pt(0)

    # Heading 2 Custom
    h2_style = styles.add_style('Heading 2 Custom', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.base_style = styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Times New Roman'
    h2_font.size = Pt(13)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(0, 0, 0)
    h2_format = h2_style.paragraph_format
    h2_format.space_before = Pt(6)
    h2_format.space_after = Pt(12)
    h2_format.line_spacing = 1.0
    h2_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2_format.left_indent = Pt(0)

    # Heading 3 Custom
    h3_style = styles.add_style('Heading 3 Custom', WD_STYLE_TYPE.PARAGRAPH)
    h3_style.base_style = styles['Heading 3']
    h3_font = h3_style.font
    h3_font.name = 'Times New Roman'
    h3_font.size = Pt(13)
    h3_font.bold = True
    h3_font.italic = True
    h3_font.color.rgb = RGBColor(0, 0, 0)
    h3_format = h3_style.paragraph_format
    h3_format.space_before = Pt(6)
    h3_format.space_after = Pt(12)
    h3_format.line_spacing = 1.0
    h3_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3_format.left_indent = Pt(0)

    # Content Custom
    content_style = styles.add_style('Content Custom', WD_STYLE_TYPE.PARAGRAPH)
    c_font = content_style.font
    c_font.name = 'Times New Roman'
    c_font.size = Pt(13)
    c_format = content_style.paragraph_format
    c_format.space_before = Pt(10)
    c_format.space_after = Pt(0)
    c_format.line_spacing = 1.5
    c_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    c_format.left_indent = Pt(0)
    c_format.first_line_indent = Pt(0)

def main():
    doc = Document()
    setup_styles(doc)

    # Content structure
    add_heading_1(doc, "CHƯƠNG 1. GIỚI THIỆU VÀ BÀI TOÁN")
    add_heading_2(doc, "1.1 Giới thiệu bài toán")
    add_content(doc, "Thị trường chứng khoán Việt Nam tạo ra một lượng lớn dữ liệu mỗi ngày, từ giao dịch, giá cả, đến các chỉ báo tài chính phức tạp. Tuy nhiên, phần lớn các công cụ báo cáo tài chính hiện tại không cung cấp dữ liệu sạch một cách tự động với đầy đủ các tầng kiểm soát chất lượng (Bronze, Silver, Gold).")
    add_content(doc, "Dự án Vietnam Stock Market Data Engineering Pipeline được xây dựng nhằm mục đích cung cấp một luồng dữ liệu (Data Pipeline) hoàn toàn tự động cho thị trường chứng khoán Việt Nam, ứng dụng kiến trúc Medallion hiện đại. Hệ thống thu thập, làm sạch, và tính toán các chỉ báo tài chính (MACD, RSI, EMA, Bollinger Bands) hoàn toàn bằng SQL/dbt để trình bày trực quan trên Power BI.")
    
    add_heading_2(doc, "1.2 Mục tiêu đề tài")
    add_content(doc, "Thu thập dữ liệu chứng khoán cuối ngày (OHLCV) từ các API công khai (như vnstock/TCBS/VCI).")
    add_content(doc, "Tự động làm sạch và chuyển đổi dữ liệu thông qua kiến trúc Medallion (Bronze, Silver, Gold).")
    add_content(doc, "Tự động tính toán các chỉ báo kỹ thuật chuyên sâu (EMA, MACD, RSI, Bollinger Bands) ở tầng Gold.")
    add_content(doc, "Điều phối hệ thống (Orchestration) bằng Apache Airflow, chạy tự động trên Docker, và trình bày trực quan (Visualization) trên Power BI.")

    add_heading_1(doc, "CHƯƠNG 2. KIẾN TRÚC HỆ THỐNG")
    add_heading_2(doc, "2.1 Tổng quan kiến trúc")
    add_content(doc, "Kiến trúc hệ thống tuân theo chuẩn Medallion, kết hợp Airflow, dbt và PostgreSQL. Luồng dữ liệu hoạt động như sau:")
    add_content(doc, "1. Provider Layer: Hợp nhất module lấy dữ liệu (VnstockProvider / MockProvider) thành một giao diện đồng nhất nhằm tận dụng khả năng tự fallback giữa TCBS/MSN của Vnstock 4.x.")
    add_content(doc, "2. Ingestion Layer: Script Python tải dữ liệu từ API và Upsert vào Data Warehouse. Logic retry và idempotent được quản lý chặt chẽ.")
    add_content(doc, "3. Bronze Layer: Dữ liệu thô lưu ở dạng JSONB trong PostgreSQL với primary key (code, date).")
    add_content(doc, "4. Silver Layer: Dùng dbt để parse JSONB, loại bỏ record xấu, thêm is_valid flag (vd: close > 0, high >= low, volume >= 0) và check data quality.")
    add_content(doc, "5. Gold Layer: Dùng dbt để chuẩn hóa thành Star Schema gồm Fact và Dimension, tính toán incremental các chỉ báo như EMA, MACD, RSI, Bollinger.")
    add_content(doc, "6. Visualization Layer: Power BI kết nối trực tiếp đến các bảng Gold (DirectQuery hoặc Import).")
    
    add_heading_2(doc, "2.2 Các quyết định kiến trúc quan trọng (ADRs)")
    add_heading_3(doc, "2.2.1 Sử dụng PostgreSQL làm Kho dữ liệu")
    add_content(doc, "PostgreSQL 17 được chọn làm Data Warehouse vì sự hỗ trợ mạnh mẽ cho JSONB (quan trọng với lớp Bronze), khả năng phân mảnh (partition), và tích hợp dbt tốt qua dbt-postgres.")
    add_heading_3(doc, "2.2.2 Sử dụng dbt Core 1.10")
    add_content(doc, "Sử dụng dbt Core 1.10 thay vì 2.0 để đảm bảo tính ổn định tối đa cho các macro và tích hợp Airflow trong giai đoạn hiện tại.")
    add_heading_3(doc, "2.2.3 EMA9 cho MACD Signal")
    add_content(doc, "Chỉ báo MACD được cấu thành từ MACD Line và MACD Signal. Thay vì dùng Simple Moving Average (SMA) cho Signal line (có rủi ro lệch chuẩn lớn), hệ thống thực hiện tính Exponential Moving Average (EMA9) chính xác cho Signal Line, đảm bảo đáp ứng chuẩn Data Quality G-03 (<0.5% sai số).")

    add_heading_1(doc, "CHƯƠNG 3. THIẾT KẾ CHI TIẾT")
    add_heading_2(doc, "3.1 Data Contracts")
    add_content(doc, "Hệ thống áp dụng hợp đồng dữ liệu nghiêm ngặt giữa các tầng.")
    add_heading_3(doc, "3.1.1 Bronze Contract")
    add_content(doc, "Lưu ở bảng bronze_prices với các schema cơ bản (code, date, open, high, low, close, volume, raw_json, source, ingested_at).")
    add_heading_3(doc, "3.1.2 Silver Contract")
    add_content(doc, "Bảng silver_stock_prices kế thừa từ Bronze, trích xuất dữ liệu, định dạng kiểu dữ liệu đúng, đánh cờ is_valid (True/False) và lý do (dq_flag) nếu giá trị đóng cửa <= 0 hoặc high < low.")
    add_heading_3(doc, "3.1.3 Gold Contract")
    add_content(doc, "Mô hình Star Schema với các bảng: fact_stock_price, fact_stock_indicators (vận hành dạng incremental 60-day lookback để tối ưu tài nguyên), fact_market_summary, dim_stock, dim_date.")

    add_heading_2(doc, "3.2 Điều phối với Apache Airflow")
    add_content(doc, "Toàn bộ chu trình được chạy dưới dạng một DAG hàng ngày (dag_daily). Mỗi task được chia nhỏ và chạy LocalExecutor.")
    add_content(doc, "Airflow trigger dbt trực tiếp qua bash scripts, sử dụng template {{ ds }} đảm bảo quá trình xử lý luôn idempotent dù backfill hay trigger bằng tay.")

    add_heading_1(doc, "CHƯƠNG 4. KẾT QUẢ VÀ ĐÁNH GIÁ")
    add_heading_2(doc, "4.1 Kết quả đạt được")
    add_content(doc, "Hệ thống Ingestion có thể tự động bypass hoặc fallback khi các API chứng khoán (VNDirect, TCBS) quá tải (rate limit), bảo đảm luồng dữ liệu không đứt gãy.")
    add_content(doc, "Cơ chế Incremental load tại tầng Gold qua dbt (delete + insert lookback) tiết kiệm đến hơn 80% khối lượng tính toán mỗi ngày, trong khi vẫn bảo lưu tính idempotent 100%.")
    add_content(doc, "Macro SQL chuyên sâu (EMA, MACD, RSI) chứng minh sự chính xác bằng kết quả kiểm thử tương đương thư viện tính toán bằng Python.")

    add_heading_2(doc, "4.2 Hạn chế và hướng phát triển")
    add_content(doc, "Hạn chế hiện tại là chưa mở rộng tính toán intraday do giới hạn free API, và PostgreSQL có thể trở thành điểm nghẽn nếu query phân tích khối lượng vượt quá ngưỡng hàng trăm triệu dòng trong vài giây.")
    add_content(doc, "Hướng phát triển: Tích hợp Kafka/Spark để tính toán streaming thời gian thực, và chuyển dịch lên Snowflake/BigQuery khi nhu cầu scale lên mức độ Enterprise.")

    add_heading_1(doc, "CHƯƠNG 5. KẾT LUẬN")
    add_content(doc, "Đồ án đã xây dựng thành công một Data Pipeline end-to-end cho thị trường chứng khoán Việt Nam theo chuẩn công nghiệp (Medallion Architecture). Tích hợp toàn diện giữa Ingestion Python hiện đại, Airflow điều phối, Data Warehouse PostgreSQL, Transformation dbt tiên tiến, và Visualization Power BI. Luồng dữ liệu đã đáp ứng độ chính xác, tính ổn định và sẵn sàng ứng dụng cho mục đích phân tích báo cáo tài chính thực tế.")

    doc.save('Bao_Cao_Do_An.docx')

if __name__ == '__main__':
    main()
